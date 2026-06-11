from __future__ import annotations

import csv
from collections import defaultdict
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from functools import lru_cache
import hashlib
from html import escape
from io import BytesIO
import json
import os
from pathlib import Path
import re
import shutil
import time as time_module
from urllib.parse import urlparse
from zoneinfo import ZoneInfo


BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "historial_registros.csv"
STATE_FILE = BASE_DIR / "estados_dia.csv"
EVIDENCE_FILE = BASE_DIR / "evidencias_registros.csv"
BACKUP_DIR = BASE_DIR / "copias_seguridad"
EVIDENCE_DIR = BASE_DIR / "evidencias"
DATE_FORMAT = "%Y-%m-%d %H:%M:%S"
COLOMBIA_TZ = ZoneInfo("America/Bogota")
TODOS_LOS_PERIODOS = "Todos los periodos"
USUARIO_PREDETERMINADO = "Andres"
COLUMNAS_ARCHIVO = ["fecha", "hora", "tipo", "detalle", "periodo", "turno", "jornada"]
COLUMNAS_ESTADOS = ["fecha", "estado", "detalle", "periodo"]
COLUMNAS_EVIDENCIAS = [
    "registro_clave",
    "fecha",
    "hora",
    "tipo",
    "detalle",
    "periodo",
    "turno",
    "jornada",
    "ubicacion_texto",
    "foto_nombre",
    "foto_url",
]
ESTADOS_DIA = [
    "12h dia",
    "12h noche",
    "5h manana",
    "6h manana",
    "6h tarde",
    "libre",
    "libre despues de noche",
    "sin definir",
]
TURNOS = {
    "12h dia": {
        "duracion_minutos": 12 * 60,
        "salida_permitida_minutos": 60,
        "inicio": (7, 0),
        "fin": (19, 0),
    },
    "12h noche": {
        "duracion_minutos": 12 * 60,
        "salida_permitida_minutos": 60,
        "inicio": (19, 0),
        "fin": (7, 0),
    },
    "5h manana": {
        "duracion_minutos": 5 * 60,
        "salida_permitida_minutos": 0,
        "inicio": (7, 0),
        "fin": (12, 0),
    },
    "6h manana": {
        "duracion_minutos": 6 * 60,
        "salida_permitida_minutos": 0,
        "inicio": (7, 0),
        "fin": (13, 0),
    },
    "6h tarde": {
        "duracion_minutos": 6 * 60,
        "salida_permitida_minutos": 0,
        "inicio": (13, 0),
        "fin": (19, 0),
    },
    "libre": {
        "duracion_minutos": 0,
        "salida_permitida_minutos": 0,
        "inicio": None,
        "fin": None,
    },
    "sin definir": {
        "duracion_minutos": 0,
        "salida_permitida_minutos": 0,
        "inicio": None,
        "fin": None,
    },
}
ESTADOS_CON_TURNO = {"12h dia", "12h noche", "5h manana", "6h manana", "6h tarde"}

tk = None
filedialog = None
messagebox = None
ttk = None
GOOGLE_WORKSHEETS_INICIALIZADAS: set[tuple[str, str]] = set()
GOOGLE_WORKSHEET_CACHE: dict[tuple[str, str], object] = {}
MESES_ES = [
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]
DIAS_SEMANA_ABREV_ES = ["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]
PERFILES_DISPONIBILIDAD = {
    "Andres": {
        "ciudad": "Sogamoso",
        "destinatarios": [
            ("Doctores", False),
            ("CAROLINA VARGAS JAIMES", True),
            ("Líder de Pediatría", False),
            ("ZULMA CRISTINA MONTAÑA MARTÍNEZ", True),
            ("Subgerente Científica", False),
            ("HOSPITAL REGIONAL DE SOGAMOSO E.S.E.", True),
        ],
        "saludo": "Cordial saludo,",
        "introduccion": (
            "Por medio del presente me permito presentar mi disponibilidad para el mes de "
            "{mes} del presente año de la siguiente manera:"
        ),
        "cierre": "Quedo atento a los comentarios.",
        "despedida": "Cordialmente,",
        "firma": [
            ("ANDRÉS ROBERTO LÓPEZ RUIZ", True),
            ("Médico Especialista en Pediatría", False),
            ("Contratista", False),
            ("HOSPITAL REGIONAL DE SOGAMOSO E.S.E.", True),
        ],
    }
    ,
    "Esposa": {
        "ciudad": "Sogamoso",
        "destinatarios": [
            ("Doctora", False),
            ("ZULMA CRISTINA MONTAÑA MARTÍNEZ", True),
            ("Subgerente Científica", False),
            ("HOSPITAL REGIONAL DE SOGAMOSO E.S.E.", True),
        ],
        "saludo": "Cordial saludo,",
        "introduccion": (
            "Por medio del presente me permito presentar mi disponibilidad para el mes de "
            "{mes} del presente año de la siguiente manera:"
        ),
        "cierre": "Quedo atento a los comentarios.",
        "despedida": "Cordialmente,",
        "firma": [
            ("LINA MARIA OSORIO REYES", True),
            ("Médico Especialista en Neonatología", False),
            ("Contratista", False),
            ("HOSPITAL REGIONAL DE SOGAMOSO E.S.E.", True),
        ],
    },
}


def usuario_actual() -> str:
    return (os.environ.get("APP_USER", USUARIO_PREDETERMINADO) or USUARIO_PREDETERMINADO).strip()


def perfil_disponibilidad_actual() -> dict | None:
    return PERFILES_DISPONIBILIDAD.get(usuario_actual())


def _slug_usuario(usuario: str | None = None) -> str:
    texto = (usuario or usuario_actual()).strip().lower()
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_") or "usuario"


def _es_usuario_predeterminado(usuario: str | None = None) -> bool:
    return (usuario or usuario_actual()).strip() == USUARIO_PREDETERMINADO


def ruta_archivo_usuario(base: Path, usuario: str | None = None) -> Path:
    if _es_usuario_predeterminado(usuario):
        return base
    return base.with_name(f"{base.stem}__{_slug_usuario(usuario)}{base.suffix}")


def carpeta_evidencias_usuario(usuario: str | None = None) -> Path:
    if _es_usuario_predeterminado(usuario):
        return EVIDENCE_DIR
    return EVIDENCE_DIR / _slug_usuario(usuario)


def nombre_worksheet_usuario(base: str, usuario: str | None = None) -> str:
    if _es_usuario_predeterminado(usuario):
        return base
    return f"{base}__{_slug_usuario(usuario)}"


def cargar_ui_escritorio():
    global tk, filedialog, messagebox, ttk
    if tk is not None:
        return

    import tkinter as tkinter_mod
    from tkinter import filedialog as filedialog_mod
    from tkinter import messagebox as messagebox_mod
    from tkinter import ttk as ttk_mod

    tk = tkinter_mod
    filedialog = filedialog_mod
    messagebox = messagebox_mod
    ttk = ttk_mod


@dataclass
class Registro:
    fecha: str
    hora: str
    tipo: str
    detalle: str
    periodo: str
    turno: str
    jornada: str

    @property
    def fecha_hora(self) -> datetime:
        return parsear_fecha_hora_registro(self.fecha, self.hora)


@dataclass
class ResumenJornada:
    jornada: str
    periodo: str
    turno: str
    horario: str
    minutos_programados: int
    minutos_dentro: int
    minutos_fuera: int
    minutos_permitidos: int
    minutos_exceso: int
    estado: str


@dataclass
class EstadoDia:
    fecha: str
    estado: str
    detalle: str
    periodo: str

    @property
    def fecha_base(self) -> date:
        return datetime.strptime(self.fecha, "%Y-%m-%d").date()


@dataclass
class EvidenciaRegistro:
    registro_clave: str
    fecha: str
    hora: str
    tipo: str
    detalle: str
    periodo: str
    turno: str
    jornada: str
    ubicacion_texto: str
    foto_nombre: str
    foto_url: str


def ahora_colombia() -> datetime:
    return datetime.now(COLOMBIA_TZ)


def turno_actual_por_hora() -> str:
    hora = ahora_colombia().hour
    if hora >= 19 or hora < 7:
        return "12h noche"
    return "12h dia"


def calcular_periodo(fecha_hora: datetime) -> str:
    if fecha_hora.day >= 21:
        inicio = fecha_hora.replace(day=21)
        if fecha_hora.month == 12:
            fin = fecha_hora.replace(year=fecha_hora.year + 1, month=1, day=20)
        else:
            fin = fecha_hora.replace(month=fecha_hora.month + 1, day=20)
    else:
        fin = fecha_hora.replace(day=20)
        if fecha_hora.month == 1:
            inicio = fecha_hora.replace(year=fecha_hora.year - 1, month=12, day=21)
        else:
            inicio = fecha_hora.replace(month=fecha_hora.month - 1, day=21)

    return f"{inicio.strftime('%Y-%m-%d')} al {fin.strftime('%Y-%m-%d')}"


def calcular_jornada(turno: str, fecha_hora: datetime) -> str:
    if turno == "12h noche":
        fecha_base = fecha_hora.date() if fecha_hora.hour >= 19 else (fecha_hora - timedelta(days=1)).date()
        return fecha_base.strftime("%Y-%m-%d")
    return fecha_hora.strftime("%Y-%m-%d")


def inicio_fin_jornada(turno: str, jornada: str) -> tuple[datetime | None, datetime | None]:
    config = TURNOS.get(turno, TURNOS["sin definir"])
    if config["inicio"] is None or config["fin"] is None:
        return None, None

    fecha_base = datetime.strptime(jornada, "%Y-%m-%d").replace(tzinfo=COLOMBIA_TZ)
    hora_inicio, minuto_inicio = config["inicio"]
    hora_fin, minuto_fin = config["fin"]

    inicio = fecha_base.replace(hour=hora_inicio, minute=minuto_inicio, second=0)
    fin = fecha_base.replace(hour=hora_fin, minute=minuto_fin, second=0)
    if turno == "12h noche":
        fin += timedelta(days=1)
    return inicio, fin


def formato_horario(turno: str, jornada: str) -> str:
    inicio, fin = inicio_fin_jornada(turno, jornada)
    if not inicio or not fin:
        return "-"
    return f"{inicio.strftime('%Y-%m-%d %H:%M')} a {fin.strftime('%Y-%m-%d %H:%M')}"


def minutos_a_texto(minutos: int) -> str:
    signo = "-" if minutos < 0 else ""
    minutos = abs(minutos)
    horas = minutos // 60
    resto = minutos % 60
    return f"{signo}{horas:02d}:{resto:02d}"


def formatear_hora_visible(hora: str) -> str:
    hora = hora.strip()
    if len(hora) >= 5:
        return hora[:5]
    return hora


def clave_mes_calendario(fecha_base: date) -> str:
    return fecha_base.strftime("%Y-%m")


def formatear_mes_calendario(clave_mes: str) -> str:
    anio, mes = clave_mes.split("-")
    return f"{MESES_ES[int(mes) - 1].capitalize()} {anio}"


def formatear_mes_disponibilidad(clave_mes: str) -> str:
    anio, mes = clave_mes.split("-")
    return f"{MESES_ES[int(mes) - 1]} de {anio}"


def parsear_fecha_hora_registro(fecha: str, hora: str) -> datetime:
    fecha = fecha.strip()
    hora = hora.strip()
    formatos = ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M")
    ultimo_error = None
    for formato in formatos:
        try:
            return datetime.strptime(f"{fecha} {hora}", formato).replace(tzinfo=COLOMBIA_TZ)
        except ValueError as exc:
            ultimo_error = exc
    raise ValueError(f"Fecha u hora de registro invalida: {fecha} {hora}") from ultimo_error


def validar_bloqueo_dia_libre(registros: list[Registro], tipo: str, fecha_hora: datetime):
    fecha_texto = fecha_hora.strftime("%Y-%m-%d")
    registros_dia = [registro for registro in registros if registro.fecha == fecha_texto]

    if tipo == "libre":
        if any(registro.tipo == "libre" for registro in registros_dia):
            raise ValueError("Ese día ya está marcado como libre.")
        if any(registro.tipo != "libre" for registro in registros_dia):
            raise ValueError("No puedes marcar libre un día que ya tiene entradas o salidas.")
        return

    if any(registro.tipo == "libre" for registro in registros_dia):
        raise ValueError("Ese día está bloqueado como libre y no admite más registros.")


def construir_fecha_hora_manual(fecha: date | datetime | str, hora: time | str) -> datetime:
    if isinstance(fecha, datetime):
        fecha_base = fecha.date()
    elif isinstance(fecha, date):
        fecha_base = fecha
    elif isinstance(fecha, str):
        fecha_base = datetime.strptime(fecha, "%Y-%m-%d").date()
    else:
        raise ValueError("Fecha manual inválida.")

    if isinstance(hora, time):
        hora_base = hora
    elif isinstance(hora, str):
        hora_texto = hora.strip()
        if hora_texto.isdigit() and len(hora_texto) in {3, 4}:
            hora_texto = hora_texto.zfill(4)
            hora_texto = f"{hora_texto[:2]}:{hora_texto[2:]}"

        formatos = ("%H:%M:%S", "%H:%M")
        ultimo_error = None
        for formato in formatos:
            try:
                hora_base = datetime.strptime(hora_texto, formato).time()
                break
            except ValueError as exc:
                ultimo_error = exc
        else:
            raise ValueError("La hora debe tener formato HH:MM o HHMM.") from ultimo_error
    else:
        raise ValueError("Hora manual inválida.")

    return datetime.combine(fecha_base, hora_base).replace(tzinfo=COLOMBIA_TZ)


def asegurar_archivo():
    BACKUP_DIR.mkdir(exist_ok=True)
    carpeta_evidencias_usuario().mkdir(exist_ok=True, parents=True)
    if usar_google_sheets():
        if not _hojas_google_listas():
            asegurar_hojas_google()
        return
    archivo_datos = ruta_archivo_usuario(DATA_FILE)
    if not archivo_datos.exists():
        with archivo_datos.open("w", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(COLUMNAS_ARCHIVO)
    asegurar_archivo_estados()
    normalizar_archivo_existente()


def asegurar_archivo_estados():
    if usar_google_sheets():
        if not _hojas_google_listas():
            asegurar_hojas_google()
        return
    archivo_estados = ruta_archivo_usuario(STATE_FILE)
    if not archivo_estados.exists():
        with archivo_estados.open("w", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(COLUMNAS_ESTADOS)


def asegurar_archivo_evidencias():
    if usar_google_sheets():
        if not _hojas_google_listas():
            asegurar_hojas_google()
        return
    archivo_evidencias = ruta_archivo_usuario(EVIDENCE_FILE)
    if not archivo_evidencias.exists():
        with archivo_evidencias.open("w", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(COLUMNAS_EVIDENCIAS)


def usar_google_sheets() -> bool:
    return bool(
        _normalizar_google_sheet_id(os.environ.get("GOOGLE_SHEET_ID", ""))
        and os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    )


def _normalizar_google_sheet_id(value: str) -> str:
    texto = str(value or "").strip()
    if not texto:
        return ""
    patron = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", texto)
    if patron:
        return patron.group(1)
    patron = re.search(r"^[a-zA-Z0-9-_]{20,}$", texto)
    if patron:
        return patron.group(0)
    return texto


@lru_cache(maxsize=2)
def _google_client_desde_json(credenciales_json: str):
    import gspread

    credenciales = json.loads(credenciales_json)
    return gspread.service_account_from_dict(credenciales)


def _google_client():
    return _google_client_desde_json(os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"])


@lru_cache(maxsize=4)
def _google_sheet_desde_id(sheet_id: str, credenciales_json: str):
    client = _google_client_desde_json(credenciales_json)
    return client.open_by_key(sheet_id)


def _google_sheet():
    sheet_id = _normalizar_google_sheet_id(os.environ["GOOGLE_SHEET_ID"])
    credenciales_json = os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"]
    try:
        return _google_sheet_desde_id(sheet_id, credenciales_json)
    except Exception as exc:
        raise RuntimeError(
            "No se pudo abrir la Google Sheet. Revisa el sheet id/url, "
            "que la hoja este compartida con la cuenta de servicio como Editor "
            "y que los Secrets de Streamlit sigan vigentes."
        ) from exc


def _obtener_worksheet(nombre: str, headers: list[str]):
    libro = _google_sheet()
    sheet_id = _normalizar_google_sheet_id(os.environ["GOOGLE_SHEET_ID"])
    cache_key = (sheet_id, nombre)
    if cache_key in GOOGLE_WORKSHEET_CACHE:
        return GOOGLE_WORKSHEET_CACHE[cache_key]

    try:
        hoja = libro.worksheet(nombre)
        GOOGLE_WORKSHEETS_INICIALIZADAS.add(cache_key)
        GOOGLE_WORKSHEET_CACHE[cache_key] = hoja
        return hoja
    except Exception:
        try:
            hoja = libro.add_worksheet(title=nombre, rows=1000, cols=max(len(headers), 6))
            hoja.append_row(headers)
            GOOGLE_WORKSHEETS_INICIALIZADAS.add(cache_key)
            GOOGLE_WORKSHEET_CACHE[cache_key] = hoja
            return hoja
        except Exception:
            # If another request created the worksheet first, fetch it again.
            hoja = libro.worksheet(nombre)
    GOOGLE_WORKSHEETS_INICIALIZADAS.add(cache_key)
    GOOGLE_WORKSHEET_CACHE[cache_key] = hoja
    return hoja


def asegurar_hojas_google():
    _obtener_worksheet(nombre_worksheet_usuario("movimientos"), COLUMNAS_ARCHIVO)
    _obtener_worksheet(nombre_worksheet_usuario("estados"), COLUMNAS_ESTADOS)
    _obtener_worksheet(nombre_worksheet_usuario("evidencias"), COLUMNAS_EVIDENCIAS)


def _hojas_google_listas() -> bool:
    if not usar_google_sheets():
        return False
    sheet_id = _normalizar_google_sheet_id(os.environ["GOOGLE_SHEET_ID"])
    requeridas = {(sheet_id, "movimientos"), (sheet_id, "estados"), (sheet_id, "evidencias")}
    return requeridas.issubset(GOOGLE_WORKSHEETS_INICIALIZADAS) and requeridas.issubset(set(GOOGLE_WORKSHEET_CACHE))


def _google_con_reintento(func, *args, **kwargs):
    ultimo_error = None
    for intento in range(3):
        try:
            return func(*args, **kwargs)
        except Exception as exc:
            ultimo_error = exc
            if intento == 2:
                raise
            time_module.sleep(1 + intento)
    raise ultimo_error


def _reescribir_worksheet(hoja, headers: list[str], filas: list[list[str]]):
    valores = [headers, *filas]
    _google_con_reintento(hoja.clear)
    _google_con_reintento(hoja.update, range_name="A1", values=valores)


def clave_registro(registro: Registro) -> str:
    return " | ".join(
        [
            registro.fecha,
            registro.hora,
            registro.tipo,
            registro.turno,
            registro.jornada,
            registro.detalle,
        ]
    )


def _drive_service():
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    credenciales = json.loads(os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"])
    scopes = ["https://www.googleapis.com/auth/drive"]
    credentials = service_account.Credentials.from_service_account_info(credenciales, scopes=scopes)
    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def usar_cloudinary() -> bool:
    return bool(
        os.environ.get("CLOUDINARY_CLOUD_NAME", "").strip()
        and os.environ.get("CLOUDINARY_UPLOAD_PRESET", "").strip()
    )


def cloudinary_puede_borrar() -> bool:
    return bool(
        os.environ.get("CLOUDINARY_CLOUD_NAME", "").strip()
        and os.environ.get("CLOUDINARY_API_KEY", "").strip()
        and os.environ.get("CLOUDINARY_API_SECRET", "").strip()
    )


def guardar_foto_evidencia(nombre_base: str, foto_bytes: bytes) -> tuple[str, str]:
    nombre_archivo = f"{nombre_base}.jpg"

    from PIL import Image

    imagen = Image.open(BytesIO(foto_bytes)).convert("RGB")
    imagen.thumbnail((1280, 1280))
    salida = BytesIO()
    imagen.save(salida, format="JPEG", quality=82, optimize=True)
    salida.seek(0)

    if usar_cloudinary():
        cloud_name = os.environ.get("CLOUDINARY_CLOUD_NAME", "").strip()
        upload_preset = os.environ.get("CLOUDINARY_UPLOAD_PRESET", "").strip()
        asset_folder = os.environ.get("CLOUDINARY_FOLDER", "").strip()
        try:
            import requests

            endpoint = f"https://api.cloudinary.com/v1_1/{cloud_name}/image/upload"
            data = {
                "upload_preset": upload_preset,
                "public_id": nombre_base,
            }
            if asset_folder:
                data["folder"] = asset_folder
            response = requests.post(
                endpoint,
                data=data,
                files={"file": (nombre_archivo, salida.getvalue(), "image/jpeg")},
                timeout=60,
            )
            response.raise_for_status()
            payload = response.json()
            foto_url = payload.get("secure_url") or payload.get("url")
            if not foto_url:
                raise RuntimeError("Cloudinary no devolvió URL de la imagen.")
            return nombre_archivo, foto_url
        except Exception as exc:
            raise RuntimeError(
                "No se pudo subir la foto a Cloudinary. "
                "Revisa cloudinary_cloud_name, cloudinary_upload_preset y el preset sin firma. "
                f"Detalle: {exc}"
            ) from exc

    if usar_google_sheets():
        carpeta_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "").strip()
        if not carpeta_id:
            raise RuntimeError(
                "Falta configurar Cloudinary o google_drive_folder_id en Streamlit Secrets para guardar fotos."
            )
        try:
            from googleapiclient.http import MediaIoBaseUpload

            servicio = _drive_service()
            metadata = {"name": nombre_archivo}
            metadata["parents"] = [carpeta_id]
            media = MediaIoBaseUpload(salida, mimetype="image/jpeg", resumable=False)
            archivo = (
                servicio.files()
                .create(body=metadata, media_body=media, fields="id, webViewLink")
                .execute()
            )
            try:
                servicio.permissions().create(
                    fileId=archivo["id"],
                    body={"type": "anyone", "role": "reader"},
                ).execute()
            except Exception:
                # Some Google Workspace / Drive configs block public-link permissions.
                # The file may still exist correctly inside the shared folder.
                pass
            foto_url = archivo.get("webViewLink") or f"https://drive.google.com/file/d/{archivo['id']}/view"
            return nombre_archivo, foto_url
        except Exception as exc:
            raise RuntimeError(
                "No se pudo subir la foto a Google Drive. Revisa la carpeta compartida, "
                "el google_drive_folder_id, la Google Drive API y los permisos de la cuenta de servicio. "
                f"Detalle: {exc}"
            ) from exc

    carpeta_destino = carpeta_evidencias_usuario()
    carpeta_destino.mkdir(exist_ok=True, parents=True)
    destino = carpeta_destino / nombre_archivo
    destino.write_bytes(salida.getvalue())
    return nombre_archivo, str(destino)


def _cloudinary_public_id_desde_evidencia(evidencia: EvidenciaRegistro) -> str:
    nombre_base = Path(evidencia.foto_nombre or "").stem
    if not nombre_base:
        ruta = urlparse(evidencia.foto_url or "").path
        if "/upload/" in ruta:
            despues = ruta.split("/upload/", 1)[1]
            partes = despues.split("/")
            while partes and re.fullmatch(r"v\d+", partes[0]):
                partes.pop(0)
            if partes:
                nombre_base = "/".join(partes)
                nombre_base = re.sub(r"\.[A-Za-z0-9]+$", "", nombre_base)
    carpeta = os.environ.get("CLOUDINARY_FOLDER", "").strip().strip("/")
    if nombre_base and carpeta and not nombre_base.startswith(f"{carpeta}/"):
        return f"{carpeta}/{nombre_base}"
    return nombre_base


def borrar_archivo_evidencia(evidencia: EvidenciaRegistro) -> None:
    if evidencia.foto_url and "res.cloudinary.com" in evidencia.foto_url and cloudinary_puede_borrar():
        public_id = _cloudinary_public_id_desde_evidencia(evidencia)
        if public_id:
            import requests

            cloud_name = os.environ["CLOUDINARY_CLOUD_NAME"].strip()
            api_key = os.environ["CLOUDINARY_API_KEY"].strip()
            api_secret = os.environ["CLOUDINARY_API_SECRET"].strip()
            timestamp = str(int(time_module.time()))
            firma_base = f"public_id={public_id}&timestamp={timestamp}{api_secret}"
            signature = hashlib.sha1(firma_base.encode("utf-8")).hexdigest()
            response = requests.post(
                f"https://api.cloudinary.com/v1_1/{cloud_name}/image/destroy",
                data={
                    "public_id": public_id,
                    "timestamp": timestamp,
                    "api_key": api_key,
                    "signature": signature,
                },
                timeout=30,
            )
            response.raise_for_status()
        return

    if evidencia.foto_url and evidencia.foto_url.startswith("/"):
        ruta = Path(evidencia.foto_url)
        if ruta.exists():
            ruta.unlink()


def evidencia_corresponde_a_registro(evidencia: EvidenciaRegistro, registro: Registro) -> bool:
    if evidencia.registro_clave == clave_registro(registro):
        return True
    return (
        evidencia.fecha == registro.fecha
        and evidencia.hora == registro.hora
        and evidencia.tipo == registro.tipo
        and evidencia.turno == registro.turno
        and evidencia.jornada == registro.jornada
    )


def normalizar_archivo_existente():
    archivo_datos = ruta_archivo_usuario(DATA_FILE)
    with archivo_datos.open("r", newline="", encoding="utf-8") as file:
        filas = list(csv.reader(file))

    if not filas:
        with archivo_datos.open("w", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(COLUMNAS_ARCHIVO)
        return

    encabezado = filas[0]
    if encabezado == COLUMNAS_ARCHIVO:
        return

    respaldo = BACKUP_DIR / f"historial_registros_migracion_{ahora_colombia().strftime('%Y%m%d_%H%M%S')}.csv"
    shutil.copy2(archivo_datos, respaldo)

    registros_convertidos: list[Registro] = []
    for fila in filas[1:]:
        mapa = dict(zip(encabezado, fila))
        fecha = mapa.get("fecha", "")
        hora = mapa.get("hora", "")
        tipo = mapa.get("tipo", "")
        detalle = mapa.get("detalle", "")
        turno = mapa.get("turno", "") or ("libre" if tipo == "libre" else "sin definir")

        if fecha and hora:
            fecha_hora = parsear_fecha_hora_registro(fecha, hora)
            periodo = mapa.get("periodo", "") or calcular_periodo(fecha_hora)
            jornada = mapa.get("jornada", "") or calcular_jornada(turno, fecha_hora)
        else:
            periodo = mapa.get("periodo", "")
            jornada = mapa.get("jornada", "")

        registros_convertidos.append(
            Registro(
                fecha=fecha,
                hora=hora,
                tipo=tipo,
                detalle=detalle,
                periodo=periodo,
                turno=turno,
                jornada=jornada,
            )
        )

    with archivo_datos.open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ARCHIVO)
        for registro in registros_convertidos:
            writer.writerow(
                [
                    registro.fecha,
                    registro.hora,
                    registro.tipo,
                    registro.detalle,
                    registro.periodo,
                    registro.turno,
                    registro.jornada,
                ]
            )


def guardar_registro(tipo: str, turno: str, detalle: str = "") -> Registro:
    return guardar_registro_en_fecha(tipo=tipo, fecha_hora=ahora_colombia(), turno=turno, detalle=detalle)


def guardar_registro_en_fecha(tipo: str, fecha_hora: datetime, turno: str, detalle: str = "") -> Registro:
    turno_normalizado = turno if turno in TURNOS else "sin definir"
    validar_bloqueo_dia_libre(leer_registros(), tipo, fecha_hora)
    periodo = calcular_periodo(fecha_hora)
    jornada = calcular_jornada(turno_normalizado, fecha_hora)
    registro = Registro(
        fecha=fecha_hora.strftime("%Y-%m-%d"),
        hora=fecha_hora.strftime("%H:%M"),
        tipo=tipo,
        detalle=detalle,
        periodo=periodo,
        turno=turno_normalizado,
        jornada=jornada,
    )

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("movimientos"), COLUMNAS_ARCHIVO)
        _google_con_reintento(
            hoja.append_row,
            [
                registro.fecha,
                registro.hora,
                registro.tipo,
                registro.detalle,
                registro.periodo,
                registro.turno,
                registro.jornada,
            ],
        )
        return registro

    with ruta_archivo_usuario(DATA_FILE).open("a", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(
            [
                registro.fecha,
                registro.hora,
                registro.tipo,
                registro.detalle,
                registro.periodo,
                registro.turno,
                registro.jornada,
            ]
        )

    crear_copia_seguridad()
    return registro


def actualizar_registro(
    registro_original: Registro,
    fecha_hora: datetime,
    tipo: str,
    turno: str,
    detalle: str = "",
) -> Registro:
    registros = leer_registros()
    restantes: list[Registro] = []
    encontrado = False
    for registro in registros:
        if not encontrado and registro == registro_original:
            encontrado = True
            continue
        restantes.append(registro)

    if not encontrado:
        raise ValueError("No se encontró el movimiento original para editar.")

    turno_normalizado = turno if turno in TURNOS else "sin definir"
    validar_bloqueo_dia_libre(restantes, tipo, fecha_hora)
    periodo = calcular_periodo(fecha_hora)
    jornada = calcular_jornada(turno_normalizado, fecha_hora)
    registro_actualizado = Registro(
        fecha=fecha_hora.strftime("%Y-%m-%d"),
        hora=fecha_hora.strftime("%H:%M"),
        tipo=tipo,
        detalle=detalle,
        periodo=periodo,
        turno=turno_normalizado,
        jornada=jornada,
    )

    restantes.append(registro_actualizado)
    restantes.sort(key=lambda item: item.fecha_hora)

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("movimientos"), COLUMNAS_ARCHIVO)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_ARCHIVO,
            [
                [
                    registro.fecha,
                    registro.hora,
                    registro.tipo,
                    registro.detalle,
                    registro.periodo,
                    registro.turno,
                    registro.jornada,
                ]
                for registro in restantes
            ],
        )
    else:
        with ruta_archivo_usuario(DATA_FILE).open("w", newline="", encoding="utf-8") as file:
            writer = csv.writer(file)
            writer.writerow(COLUMNAS_ARCHIVO)
            for registro in restantes:
                writer.writerow(
                    [
                        registro.fecha,
                        registro.hora,
                        registro.tipo,
                        registro.detalle,
                        registro.periodo,
                        registro.turno,
                        registro.jornada,
                    ]
                )
        crear_copia_seguridad()

    actualizar_evidencias_por_registro(registro_original, registro_actualizado)
    return registro_actualizado


def eliminar_registro(registro_objetivo: Registro) -> bool:
    registros = leer_registros()
    eliminado = False
    restantes: list[Registro] = []

    for registro in registros:
        if not eliminado and registro == registro_objetivo:
            eliminado = True
            continue
        restantes.append(registro)

    if not eliminado:
        return False

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("movimientos"), COLUMNAS_ARCHIVO)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_ARCHIVO,
            [
                [
                    registro.fecha,
                    registro.hora,
                    registro.tipo,
                    registro.detalle,
                    registro.periodo,
                    registro.turno,
                    registro.jornada,
                ]
                for registro in restantes
            ],
        )
        eliminar_evidencias_por_registro(registro_objetivo)
        return True

    with ruta_archivo_usuario(DATA_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ARCHIVO)
        for registro in restantes:
            writer.writerow(
                [
                    registro.fecha,
                    registro.hora,
                    registro.tipo,
                    registro.detalle,
                    registro.periodo,
                    registro.turno,
                    registro.jornada,
                ]
            )

    eliminar_evidencias_por_registro(registro_objetivo)
    crear_copia_seguridad()
    return True


def leer_registros() -> list[Registro]:
    asegurar_archivo()
    registros: list[Registro] = []

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("movimientos"), COLUMNAS_ARCHIVO)
        filas = _google_con_reintento(hoja.get_all_records)
        for row in filas:
            fecha = str(row.get("fecha", "")).strip()
            hora = str(row.get("hora", "")).strip()
            tipo = str(row.get("tipo", "")).strip()
            if not fecha or not hora or not tipo:
                continue
            turno = str(row.get("turno", "")).strip() or ("libre" if tipo == "libre" else "sin definir")
            fecha_hora = parsear_fecha_hora_registro(fecha, hora)
            registros.append(
                Registro(
                    fecha=fecha,
                    hora=hora,
                    tipo=tipo,
                    detalle=str(row.get("detalle", "")).strip(),
                    periodo=str(row.get("periodo", "")).strip() or calcular_periodo(fecha_hora),
                    turno=turno,
                    jornada=str(row.get("jornada", "")).strip() or calcular_jornada(turno, fecha_hora),
                )
            )
        registros.sort(key=lambda item: item.fecha_hora)
        return registros

    with ruta_archivo_usuario(DATA_FILE).open("r", newline="", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            fecha = row.get("fecha", "")
            hora = row.get("hora", "")
            tipo = row.get("tipo", "")
            turno = row.get("turno", "") or ("libre" if tipo == "libre" else "sin definir")
            fecha_hora = None
            if fecha and hora:
                fecha_hora = parsear_fecha_hora_registro(fecha, hora)

            registros.append(
                Registro(
                    fecha=fecha,
                    hora=hora,
                    tipo=tipo,
                    detalle=row.get("detalle", ""),
                    periodo=row.get("periodo", "") or (calcular_periodo(fecha_hora) if fecha_hora else ""),
                    turno=turno,
                    jornada=row.get("jornada", "") or (calcular_jornada(turno, fecha_hora) if fecha_hora else ""),
                )
            )

    registros.sort(key=lambda item: item.fecha_hora)
    return registros


def guardar_evidencia_registro(
    registro: Registro,
    ubicacion_texto: str = "",
    foto_bytes: bytes | None = None,
) -> EvidenciaRegistro | None:
    ubicacion_texto = (ubicacion_texto or "").strip()
    foto_nombre = ""
    foto_url = ""

    if not ubicacion_texto and not foto_bytes:
        return None

    if foto_bytes:
        marca_unica = ahora_colombia().strftime("%Y%m%d_%H%M%S_%f")
        base_nombre = (
            f"evidencia_{_slug_usuario()}_{registro.fecha}_{registro.hora.replace(':', '')}_{registro.tipo}_{registro.turno}_{marca_unica}"
        )
        foto_nombre, foto_url = guardar_foto_evidencia(base_nombre, foto_bytes)

    evidencia = EvidenciaRegistro(
        registro_clave=clave_registro(registro),
        fecha=registro.fecha,
        hora=registro.hora,
        tipo=registro.tipo,
        detalle=registro.detalle,
        periodo=registro.periodo,
        turno=registro.turno,
        jornada=registro.jornada,
        ubicacion_texto=ubicacion_texto,
        foto_nombre=foto_nombre,
        foto_url=foto_url,
    )

    asegurar_archivo_evidencias()
    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("evidencias"), COLUMNAS_EVIDENCIAS)
        _google_con_reintento(
            hoja.append_row,
            [
                evidencia.registro_clave,
                evidencia.fecha,
                evidencia.hora,
                evidencia.tipo,
                evidencia.detalle,
                evidencia.periodo,
                evidencia.turno,
                evidencia.jornada,
                evidencia.ubicacion_texto,
                evidencia.foto_nombre,
                evidencia.foto_url,
            ],
        )
        return evidencia

    with ruta_archivo_usuario(EVIDENCE_FILE).open("a", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(
            [
                evidencia.registro_clave,
                evidencia.fecha,
                evidencia.hora,
                evidencia.tipo,
                evidencia.detalle,
                evidencia.periodo,
                evidencia.turno,
                evidencia.jornada,
                evidencia.ubicacion_texto,
                evidencia.foto_nombre,
                evidencia.foto_url,
            ]
        )
    return evidencia


def leer_evidencias() -> list[EvidenciaRegistro]:
    asegurar_archivo_evidencias()
    evidencias: list[EvidenciaRegistro] = []

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("evidencias"), COLUMNAS_EVIDENCIAS)
        filas = _google_con_reintento(hoja.get_all_records)
        for row in filas:
            clave = str(row.get("registro_clave", "")).strip()
            if not clave:
                continue
            evidencias.append(
                EvidenciaRegistro(
                    registro_clave=clave,
                    fecha=str(row.get("fecha", "")).strip(),
                    hora=str(row.get("hora", "")).strip(),
                    tipo=str(row.get("tipo", "")).strip(),
                    detalle=str(row.get("detalle", "")).strip(),
                    periodo=str(row.get("periodo", "")).strip(),
                    turno=str(row.get("turno", "")).strip(),
                    jornada=str(row.get("jornada", "")).strip(),
                    ubicacion_texto=str(row.get("ubicacion_texto", "")).strip(),
                    foto_nombre=str(row.get("foto_nombre", "")).strip(),
                    foto_url=str(row.get("foto_url", "")).strip(),
                )
            )
        return evidencias

    with ruta_archivo_usuario(EVIDENCE_FILE).open("r", newline="", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            clave = row.get("registro_clave", "").strip()
            if not clave:
                continue
            evidencias.append(
                EvidenciaRegistro(
                    registro_clave=clave,
                    fecha=row.get("fecha", "").strip(),
                    hora=row.get("hora", "").strip(),
                    tipo=row.get("tipo", "").strip(),
                    detalle=row.get("detalle", "").strip(),
                    periodo=row.get("periodo", "").strip(),
                    turno=row.get("turno", "").strip(),
                    jornada=row.get("jornada", "").strip(),
                    ubicacion_texto=row.get("ubicacion_texto", "").strip(),
                    foto_nombre=row.get("foto_nombre", "").strip(),
                    foto_url=row.get("foto_url", "").strip(),
                )
            )
    return evidencias


def actualizar_evidencias_por_registro(registro_original: Registro, registro_actualizado: Registro) -> int:
    evidencias = leer_evidencias()
    cambios = 0
    actualizadas: list[EvidenciaRegistro] = []

    for evidencia in evidencias:
        if evidencia_corresponde_a_registro(evidencia, registro_original):
            cambios += 1
            actualizadas.append(
                EvidenciaRegistro(
                    registro_clave=clave_registro(registro_actualizado),
                    fecha=registro_actualizado.fecha,
                    hora=registro_actualizado.hora,
                    tipo=registro_actualizado.tipo,
                    detalle=registro_actualizado.detalle,
                    periodo=registro_actualizado.periodo,
                    turno=registro_actualizado.turno,
                    jornada=registro_actualizado.jornada,
                    ubicacion_texto=evidencia.ubicacion_texto,
                    foto_nombre=evidencia.foto_nombre,
                    foto_url=evidencia.foto_url,
                )
            )
        else:
            actualizadas.append(evidencia)

    if cambios == 0:
        return 0

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("evidencias"), COLUMNAS_EVIDENCIAS)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_EVIDENCIAS,
            [
                [
                    evidencia.registro_clave,
                    evidencia.fecha,
                    evidencia.hora,
                    evidencia.tipo,
                    evidencia.detalle,
                    evidencia.periodo,
                    evidencia.turno,
                    evidencia.jornada,
                    evidencia.ubicacion_texto,
                    evidencia.foto_nombre,
                    evidencia.foto_url,
                ]
                for evidencia in actualizadas
            ],
        )
        return cambios

    with ruta_archivo_usuario(EVIDENCE_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_EVIDENCIAS)
        for evidencia in actualizadas:
            writer.writerow(
                [
                    evidencia.registro_clave,
                    evidencia.fecha,
                    evidencia.hora,
                    evidencia.tipo,
                    evidencia.detalle,
                    evidencia.periodo,
                    evidencia.turno,
                    evidencia.jornada,
                    evidencia.ubicacion_texto,
                    evidencia.foto_nombre,
                    evidencia.foto_url,
                ]
            )
    return cambios


def eliminar_evidencias_por_registro(registro: Registro) -> int:
    if not registro:
        return 0

    evidencias = leer_evidencias()
    a_eliminar = [evidencia for evidencia in evidencias if evidencia_corresponde_a_registro(evidencia, registro)]
    restantes = [evidencia for evidencia in evidencias if not evidencia_corresponde_a_registro(evidencia, registro)]
    eliminadas = len(evidencias) - len(restantes)

    if eliminadas == 0:
        return 0

    for evidencia in a_eliminar:
        try:
            borrar_archivo_evidencia(evidencia)
        except Exception:
            # Keep deleting the registry rows even if the external file cleanup fails.
            pass

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("evidencias"), COLUMNAS_EVIDENCIAS)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_EVIDENCIAS,
            [
                [
                    evidencia.registro_clave,
                    evidencia.fecha,
                    evidencia.hora,
                    evidencia.tipo,
                    evidencia.detalle,
                    evidencia.periodo,
                    evidencia.turno,
                    evidencia.jornada,
                    evidencia.ubicacion_texto,
                    evidencia.foto_nombre,
                    evidencia.foto_url,
                ]
                for evidencia in restantes
            ],
        )
        return eliminadas

    with ruta_archivo_usuario(EVIDENCE_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_EVIDENCIAS)
        for evidencia in restantes:
            writer.writerow(
                [
                    evidencia.registro_clave,
                    evidencia.fecha,
                    evidencia.hora,
                    evidencia.tipo,
                    evidencia.detalle,
                    evidencia.periodo,
                    evidencia.turno,
                    evidencia.jornada,
                    evidencia.ubicacion_texto,
                    evidencia.foto_nombre,
                    evidencia.foto_url,
                ]
            )
    return eliminadas


def crear_copia_seguridad():
    asegurar_archivo()
    if usar_google_sheets():
        return
    fecha_respaldo = ahora_colombia().strftime("%Y%m%d")
    destino = BACKUP_DIR / f"historial_registros_{fecha_respaldo}.csv"
    shutil.copy2(ruta_archivo_usuario(DATA_FILE), destino)


def guardar_estado_dia(fecha: date | str, estado: str, detalle: str = "") -> EstadoDia:
    asegurar_archivo_estados()
    if estado not in ESTADOS_DIA:
        raise ValueError("Estado del dia invalido.")

    if isinstance(fecha, date):
        fecha_texto = fecha.strftime("%Y-%m-%d")
        fecha_base = fecha
    elif isinstance(fecha, str):
        fecha_texto = fecha.strip()
        fecha_base = datetime.strptime(fecha_texto, "%Y-%m-%d").date()
    else:
        raise ValueError("Fecha del estado invalida.")

    periodo = calcular_periodo(datetime.combine(fecha_base, time(0, 0)).replace(tzinfo=COLOMBIA_TZ))
    estado_dia = EstadoDia(fecha=fecha_texto, estado=estado, detalle=detalle.strip(), periodo=periodo)

    estados = leer_estados_dia()
    actualizados = [
        item
        for item in estados
        if not (
            item.fecha == estado_dia.fecha
            and item.estado == estado_dia.estado
            and item.detalle == estado_dia.detalle
            and item.periodo == estado_dia.periodo
        )
    ]
    actualizados.append(estado_dia)
    actualizados.sort(key=lambda item: item.fecha)

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("estados"), COLUMNAS_ESTADOS)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_ESTADOS,
            [[item.fecha, item.estado, item.detalle, item.periodo] for item in actualizados],
        )
        return estado_dia

    with ruta_archivo_usuario(STATE_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ESTADOS)
        for item in actualizados:
            writer.writerow([item.fecha, item.estado, item.detalle, item.periodo])

    return estado_dia


def actualizar_estado_dia(
    estado_original: EstadoDia,
    fecha: date | str,
    estado: str,
    detalle: str = "",
) -> EstadoDia:
    asegurar_archivo_estados()
    if estado not in ESTADOS_DIA:
        raise ValueError("Estado del dia invalido.")

    if isinstance(fecha, date):
        fecha_texto = fecha.strftime("%Y-%m-%d")
        fecha_base = fecha
    elif isinstance(fecha, str):
        fecha_texto = fecha.strip()
        fecha_base = datetime.strptime(fecha_texto, "%Y-%m-%d").date()
    else:
        raise ValueError("Fecha del estado invalida.")

    periodo = calcular_periodo(datetime.combine(fecha_base, time(0, 0)).replace(tzinfo=COLOMBIA_TZ))
    estado_actualizado = EstadoDia(
        fecha=fecha_texto,
        estado=estado,
        detalle=detalle.strip(),
        periodo=periodo,
    )

    estados = leer_estados_dia()
    actualizados: list[EstadoDia] = []
    reemplazado = False
    for item in estados:
        if (
            not reemplazado
            and item.fecha == estado_original.fecha
            and item.estado == estado_original.estado
            and item.detalle == estado_original.detalle
            and item.periodo == estado_original.periodo
        ):
            actualizados.append(estado_actualizado)
            reemplazado = True
        else:
            actualizados.append(item)

    if not reemplazado:
        raise ValueError("No se encontró el estado original para editar.")

    estados_sin_duplicados: list[EstadoDia] = []
    for item in actualizados:
        if not any(
            existente.fecha == item.fecha
            and existente.estado == item.estado
            and existente.detalle == item.detalle
            and existente.periodo == item.periodo
            for existente in estados_sin_duplicados
        ):
            estados_sin_duplicados.append(item)

    estados_sin_duplicados.sort(key=lambda item: item.fecha)

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("estados"), COLUMNAS_ESTADOS)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_ESTADOS,
            [[item.fecha, item.estado, item.detalle, item.periodo] for item in estados_sin_duplicados],
        )
        return estado_actualizado

    with ruta_archivo_usuario(STATE_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ESTADOS)
        for item in estados_sin_duplicados:
            writer.writerow([item.fecha, item.estado, item.detalle, item.periodo])

    return estado_actualizado


def eliminar_estado_dia(estado_objetivo: EstadoDia | date | str) -> bool:
    estados = leer_estados_dia()

    if isinstance(estado_objetivo, EstadoDia):
        restantes = [
            item
            for item in estados
            if not (
                item.fecha == estado_objetivo.fecha
                and item.estado == estado_objetivo.estado
                and item.detalle == estado_objetivo.detalle
                and item.periodo == estado_objetivo.periodo
            )
        ]
    else:
        if isinstance(estado_objetivo, date):
            fecha_texto = estado_objetivo.strftime("%Y-%m-%d")
        elif isinstance(estado_objetivo, str):
            fecha_texto = estado_objetivo.strip()
        else:
            return False
        restantes = [item for item in estados if item.fecha != fecha_texto]

    if len(restantes) == len(estados):
        return False

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("estados"), COLUMNAS_ESTADOS)
        _reescribir_worksheet(
            hoja,
            COLUMNAS_ESTADOS,
            [[item.fecha, item.estado, item.detalle, item.periodo] for item in restantes],
        )
        return True

    with ruta_archivo_usuario(STATE_FILE).open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ESTADOS)
        for item in restantes:
            writer.writerow([item.fecha, item.estado, item.detalle, item.periodo])

    return True


def leer_estados_dia() -> list[EstadoDia]:
    asegurar_archivo_estados()
    estados: list[EstadoDia] = []

    if usar_google_sheets():
        hoja = _obtener_worksheet(nombre_worksheet_usuario("estados"), COLUMNAS_ESTADOS)
        for row in _google_con_reintento(hoja.get_all_records):
            fecha = str(row.get("fecha", "")).strip()
            if not fecha:
                continue
            estado = str(row.get("estado", "")).strip() or "sin definir"
            detalle = str(row.get("detalle", "")).strip()
            periodo = str(row.get("periodo", "")).strip()
            if not periodo:
                fecha_base = datetime.strptime(fecha, "%Y-%m-%d").date()
                periodo = calcular_periodo(datetime.combine(fecha_base, time(0, 0)).replace(tzinfo=COLOMBIA_TZ))
            estados.append(EstadoDia(fecha=fecha, estado=estado, detalle=detalle, periodo=periodo))
        estados.sort(key=lambda item: item.fecha)
        return estados

    with ruta_archivo_usuario(STATE_FILE).open("r", newline="", encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for row in reader:
            fecha = row.get("fecha", "").strip()
            if not fecha:
                continue
            estado = row.get("estado", "").strip() or "sin definir"
            detalle = row.get("detalle", "").strip()
            periodo = row.get("periodo", "").strip()
            if not periodo:
                fecha_base = datetime.strptime(fecha, "%Y-%m-%d").date()
                periodo = calcular_periodo(datetime.combine(fecha_base, time(0, 0)).replace(tzinfo=COLOMBIA_TZ))
            estados.append(EstadoDia(fecha=fecha, estado=estado, detalle=detalle, periodo=periodo))

    estados.sort(key=lambda item: item.fecha)
    return estados


def meses_disponibilidad_mensual(estados: list[EstadoDia]) -> list[str]:
    meses = sorted(
        {
            clave_mes_calendario(item.fecha_base)
            for item in estados
            if item.estado != "sin definir"
        }
    )
    return meses


def estados_mes_calendario(estados: list[EstadoDia], clave_mes: str) -> list[EstadoDia]:
    return [
        item
        for item in estados
        if clave_mes_calendario(item.fecha_base) == clave_mes and item.estado != "sin definir"
    ]


def _rango_turno_disponibilidad(estado: str) -> str:
    rangos = {
        "12h dia": "7:00 a.m. - 7:00 p.m.",
        "12h noche": "7:00 p.m. - 7:00 a.m.",
        "5h manana": "7:00 a.m. - 12:00 p.m.",
        "6h manana": "7:00 a.m. - 1:00 p.m.",
        "6h tarde": "1:00 p.m. - 7:00 p.m.",
    }
    return rangos.get(estado, "")


def _sufijo_turno_disponibilidad(estado: str) -> str:
    sufijos = {
        "12h dia": "DÍA",
        "12h noche": "NOCHE",
        "5h manana": "MAÑANA",
        "6h manana": "MAÑANA",
        "6h tarde": "TARDE",
    }
    return sufijos.get(estado, estado.upper())


def _expandir_servicio_disponibilidad(detalle: str) -> str:
    texto = detalle.strip().upper()
    if not texto:
        return ""
    reemplazos = {
        "HX": "HOSPITALIZACIÓN",
        "HOSP": "HOSPITALIZACIÓN",
        "HOSPITALIZACION": "HOSPITALIZACIÓN",
        "URG": "URGENCIAS",
        "UCIN": "UCIN",
    }
    if texto in reemplazos:
        return reemplazos[texto]
    return texto


def descripcion_disponibilidad_estado(estado: str, detalle: str) -> str:
    detalle_limpio = detalle.strip()
    if not detalle_limpio:
        return {
            "12h dia": "TURNO DÍA",
            "12h noche": "TURNO NOCHE",
            "5h manana": "TURNO MAÑANA",
            "6h manana": "TURNO MAÑANA",
            "6h tarde": "TURNO TARDE",
        }.get(estado, estado.upper())

    texto = _expandir_servicio_disponibilidad(detalle_limpio)
    if " " in texto or any(palabra in texto for palabra in ("DÍA", "NOCHE", "MAÑANA", "TARDE")):
        return texto
    return f"{texto} {_sufijo_turno_disponibilidad(estado)}"


def lineas_disponibilidad_mensual(estados: list[EstadoDia], clave_mes: str) -> list[str]:
    candidatos = [
        item
        for item in estados_mes_calendario(estados, clave_mes)
        if item.estado in ESTADOS_CON_TURNO
    ]
    candidatos.sort(key=lambda item: (item.fecha, _orden_estado_turno(item.estado), item.detalle))

    lineas: list[str] = []
    fecha_previa = None
    for item in candidatos:
        fecha_base = item.fecha_base
        prefijo = ""
        if fecha_previa != item.fecha:
            prefijo = f"{fecha_base.day} {DIAS_SEMANA_ABREV_ES[fecha_base.weekday()]} "
        rango = _rango_turno_disponibilidad(item.estado)
        descripcion = descripcion_disponibilidad_estado(item.estado, item.detalle)
        lineas.append(f"{prefijo}{rango} {descripcion}".strip())
        fecha_previa = item.fecha
    return lineas


def nombre_archivo_disponibilidad(clave_mes: str) -> str:
    anio, mes = clave_mes.split("-")
    mes_texto = MESES_ES[int(mes) - 1]
    return f"disponibilidad_{mes_texto}_{anio}.docx"


def generar_disponibilidad_mensual_docx(estados: list[EstadoDia], clave_mes: str) -> bytes:
    perfil = perfil_disponibilidad_actual()
    if perfil is None:
        raise ValueError("Este usuario todavía no tiene una plantilla de disponibilidad configurada.")

    lineas = lineas_disponibilidad_mensual(estados, clave_mes)
    if not lineas:
        raise ValueError("No hay turnos programados en ese mes para generar la disponibilidad.")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except Exception as exc:
        raise RuntimeError("No se pudo cargar python-docx para generar el archivo Word.") from exc

    ahora = ahora_colombia()
    anio, mes = clave_mes.split("-")
    mes_texto = MESES_ES[int(mes) - 1]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = Pt(11)

    def agregar_parrafo(texto: str = "", bold: bool = False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(texto)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = bold
        return p

    agregar_parrafo(
        f"{perfil['ciudad']} {ahora.day} de {MESES_ES[ahora.month - 1]} de {ahora.year}"
    )
    agregar_parrafo()
    for texto, bold in perfil["destinatarios"]:
        agregar_parrafo(texto, bold=bold)
    agregar_parrafo()

    asunto = doc.add_paragraph()
    asunto.paragraph_format.space_before = Pt(0)
    asunto.paragraph_format.space_after = Pt(0)
    run_asunto_1 = asunto.add_run("Asunto: ")
    run_asunto_1.bold = True
    run_asunto_1.font.name = "Arial"
    run_asunto_1.font.size = Pt(11)
    run_asunto_2 = asunto.add_run(f"Disponibilidad del mes de {mes_texto} de {anio}.")
    run_asunto_2.font.name = "Arial"
    run_asunto_2.font.size = Pt(11)

    agregar_parrafo()
    agregar_parrafo(perfil["saludo"])
    agregar_parrafo()
    agregar_parrafo(perfil["introduccion"].format(mes=mes_texto))
    agregar_parrafo()

    for linea in lineas:
        agregar_parrafo(linea)

    agregar_parrafo()
    agregar_parrafo(perfil["cierre"])
    agregar_parrafo()
    agregar_parrafo(perfil["despedida"])
    agregar_parrafo()
    agregar_parrafo()
    agregar_parrafo()
    agregar_parrafo()
    for texto, bold in perfil["firma"]:
        agregar_parrafo(texto, bold=bold)

    salida = BytesIO()
    doc.save(salida)
    return salida.getvalue()


def periodos_disponibles(registros: list[Registro]) -> list[str]:
    periodos = sorted({registro.periodo for registro in registros if registro.periodo})
    return [TODOS_LOS_PERIODOS, *periodos]


def periodos_combinados(registros: list[Registro], estados: list[EstadoDia]) -> list[str]:
    periodos = sorted(
        {registro.periodo for registro in registros if registro.periodo}
        | {estado.periodo for estado in estados if estado.periodo}
    )
    return [TODOS_LOS_PERIODOS, *periodos]


def agrupar_resumenes_jornada(registros: list[Registro]) -> list[ResumenJornada]:
    grupos: dict[tuple[str, str], list[Registro]] = defaultdict(list)
    for registro in registros:
        if registro.turno not in {"12h dia", "12h noche", "5h manana", "6h manana", "6h tarde"}:
            continue
        grupos[(registro.turno, registro.jornada)].append(registro)

    resumenes: list[ResumenJornada] = []
    for (turno, jornada), items in sorted(grupos.items(), key=lambda dato: dato[1][0].fecha_hora):
        items.sort(key=lambda item: item.fecha_hora)
        config = TURNOS[turno]
        minutos_fuera = 0
        salida_abierta: datetime | None = None
        incidencias: list[str] = []
        inicio_jornada, fin_jornada = inicio_fin_jornada(turno, jornada)

        if items and items[0].tipo != "entrada":
            incidencias.append("revisar primera marca")

        for item in items:
            momento = item.fecha_hora
            if item.tipo == "salida":
                if salida_abierta is None:
                    salida_abierta = momento
                else:
                    incidencias.append("salidas consecutivas")
                    salida_abierta = momento
            elif item.tipo == "entrada":
                if salida_abierta is not None:
                    inicio_calculo = max(salida_abierta, inicio_jornada) if inicio_jornada else salida_abierta
                    fin_calculo = min(momento, fin_jornada) if fin_jornada else momento
                    if fin_calculo > inicio_calculo:
                        minutos_fuera += int((fin_calculo - inicio_calculo).total_seconds() // 60)
                    salida_abierta = None

        if salida_abierta is not None:
            if fin_jornada and salida_abierta >= fin_jornada - timedelta(minutes=5):
                salida_abierta = None
            else:
                incidencias.append("ultima salida sin regreso")

        minutos_programados = config["duracion_minutos"]
        minutos_permitidos = config["salida_permitida_minutos"]
        minutos_exceso = max(0, minutos_fuera - minutos_permitidos)
        minutos_dentro = max(0, minutos_programados - minutos_fuera)
        if incidencias:
            estado = ", ".join(dict.fromkeys(incidencias))
        elif minutos_exceso > 0:
            estado = "ok (ojo)"
        else:
            estado = "ok"

        resumenes.append(
            ResumenJornada(
                jornada=jornada,
                periodo=items[0].periodo,
                turno=turno,
                horario=formato_horario(turno, jornada),
                minutos_programados=minutos_programados,
                minutos_dentro=minutos_dentro,
                minutos_fuera=minutos_fuera,
                minutos_permitidos=minutos_permitidos,
                minutos_exceso=minutos_exceso,
                estado=estado,
            )
        )

    return resumenes


def _combinar_estado_resumen(estado_base: str, avisos: list[str]) -> str:
    if not avisos and estado_base in {"ok", "ok (ojo)"}:
        return estado_base

    partes = []
    if estado_base and estado_base != "ok":
        partes.append(estado_base)
    for aviso in avisos:
        if aviso and aviso not in partes:
            partes.append(aviso)
    return "ok" if not partes else ", ".join(partes)


def _turno_programado_vencido(fecha: str, estado_programado: str, referencia: datetime | None = None) -> bool:
    if estado_programado not in ESTADOS_CON_TURNO:
        return False
    referencia = referencia or ahora_colombia()
    _, fin = inicio_fin_jornada(estado_programado, fecha)
    if fin is None:
        return False
    return referencia >= fin


def _orden_estado_turno(valor: str) -> int:
    orden = {
        "5h manana": 0,
        "6h manana": 1,
        "6h tarde": 2,
        "12h dia": 3,
        "12h noche": 4,
        "libre despues de noche": 5,
        "libre": 6,
        "sin definir": 7,
        "sin programar": 8,
    }
    return orden.get(valor, 99)


def resumenes_programado_vs_real(
    registros: list[Registro],
    estados: list[EstadoDia],
    referencia: datetime | None = None,
) -> list[tuple[ResumenJornada, str]]:
    referencia = referencia or ahora_colombia()
    resumenes_base = agrupar_resumenes_jornada(registros)
    resumenes_por_jornada: dict[str, list[ResumenJornada]] = defaultdict(list)
    for item in resumenes_base:
        resumenes_por_jornada[item.jornada].append(item)
    for jornada in resumenes_por_jornada:
        resumenes_por_jornada[jornada].sort(key=lambda item: (_orden_estado_turno(item.turno), item.horario))

    estados_por_fecha: dict[str, list[EstadoDia]] = defaultdict(list)
    for item in estados:
        estados_por_fecha[item.fecha].append(item)
    for fecha in estados_por_fecha:
        estados_por_fecha[fecha].sort(key=lambda item: (_orden_estado_turno(item.estado), item.detalle))

    jornadas = sorted(set(resumenes_por_jornada) | set(estados_por_fecha))
    filas: list[tuple[ResumenJornada, str]] = []

    for jornada in jornadas:
        resumenes_jornada = list(resumenes_por_jornada.get(jornada, []))
        estados_jornada = list(estados_por_fecha.get(jornada, []))
        usados_estados: set[int] = set()

        for resumen in resumenes_jornada:
            match_index = None
            for indice, estado_dia in enumerate(estados_jornada):
                if indice in usados_estados:
                    continue
                if estado_dia.estado == resumen.turno:
                    match_index = indice
                    break
            if match_index is None:
                for indice, estado_dia in enumerate(estados_jornada):
                    if indice not in usados_estados:
                        match_index = indice
                        break

            estado_dia = estados_jornada[match_index] if match_index is not None else None
            if match_index is not None:
                usados_estados.add(match_index)

            programado = estado_dia.estado if estado_dia else "sin programar"
            avisos: list[str] = []
            if estado_dia is None:
                avisos.append("sin programacion")
            elif estado_dia.estado in ESTADOS_CON_TURNO and resumen.turno != estado_dia.estado:
                avisos.append("turno distinto a programado")
            elif estado_dia.estado in {"libre", "libre despues de noche"}:
                avisos.append("movimientos en dia libre")

            filas.append(
                (
                    ResumenJornada(
                        jornada=resumen.jornada,
                        periodo=resumen.periodo,
                        turno=resumen.turno,
                        horario=resumen.horario,
                        minutos_programados=resumen.minutos_programados,
                        minutos_dentro=resumen.minutos_dentro,
                        minutos_fuera=resumen.minutos_fuera,
                        minutos_permitidos=resumen.minutos_permitidos,
                        minutos_exceso=resumen.minutos_exceso,
                        estado=_combinar_estado_resumen(resumen.estado, avisos),
                    ),
                    programado,
                )
            )

        for indice, estado_dia in enumerate(estados_jornada):
            if indice in usados_estados:
                continue

            if estado_dia.estado in ESTADOS_CON_TURNO:
                config = TURNOS[estado_dia.estado]
                vencido = _turno_programado_vencido(jornada, estado_dia.estado, referencia)
                filas.append(
                    (
                        ResumenJornada(
                            jornada=jornada,
                            periodo=estado_dia.periodo,
                            turno="sin registros" if vencido else estado_dia.estado,
                            horario=formato_horario(estado_dia.estado, jornada),
                            minutos_programados=config["duracion_minutos"],
                            minutos_dentro=0,
                            minutos_fuera=0,
                            minutos_permitidos=config["salida_permitida_minutos"],
                            minutos_exceso=0,
                            estado="sin registros para turno programado" if vencido else "programado",
                        ),
                        estado_dia.estado,
                    )
                )
            elif estado_dia.estado in {"libre", "libre despues de noche"}:
                filas.append(
                    (
                        ResumenJornada(
                            jornada=jornada,
                            periodo=estado_dia.periodo,
                            turno="sin registros",
                            horario="Libre",
                            minutos_programados=0,
                            minutos_dentro=0,
                            minutos_fuera=0,
                            minutos_permitidos=0,
                            minutos_exceso=0,
                            estado="libre",
                        ),
                        estado_dia.estado,
                    )
                )

    return filas


def resumir_periodo(registros: list[Registro]) -> dict[str, int]:
    resumenes = agrupar_resumenes_jornada(registros)
    return {
        "turnos": len(resumenes),
        "programados": sum(item.minutos_programados for item in resumenes),
        "dentro": sum(item.minutos_dentro for item in resumenes),
        "fuera": sum(item.minutos_fuera for item in resumenes),
        "permitidos": sum(item.minutos_permitidos for item in resumenes),
        "exceso": sum(item.minutos_exceso for item in resumenes),
        "sin_turno": len([item for item in registros if item.turno == "sin definir"]),
    }


def resumen_total_jornadas(resumenes: list[ResumenJornada]) -> ResumenJornada:
    total_programado = sum(item.minutos_programados for item in resumenes)
    total_dentro = sum(item.minutos_dentro for item in resumenes)
    total_fuera = sum(item.minutos_fuera for item in resumenes)
    total_permitido = sum(item.minutos_permitidos for item in resumenes)
    total_exceso = sum(item.minutos_exceso for item in resumenes)
    turnos_reales = sum(1 for item in resumenes if item.minutos_programados > 0)
    estado_total = "ok" if total_fuera <= total_permitido else "revisar"

    return ResumenJornada(
        jornada="TOTAL",
        periodo=resumenes[0].periodo if resumenes else "",
        turno=f"{turnos_reales} turnos",
        horario=f"Acumulado trabajado: {minutos_a_texto(total_dentro)}",
        minutos_programados=total_programado,
        minutos_dentro=total_dentro,
        minutos_fuera=total_fuera,
        minutos_permitidos=total_permitido,
        minutos_exceso=total_exceso,
        estado=estado_total,
    )


def exportar_csv(destino: Path, registros: list[Registro]):
    with destino.open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(COLUMNAS_ARCHIVO)
        for registro in registros:
            writer.writerow(
                [
                    registro.fecha,
                    registro.hora,
                    registro.tipo,
                    registro.detalle,
                    registro.periodo,
                    registro.turno,
                    registro.jornada,
                ]
            )


def exportar_html(destino: Path, registros: list[Registro]):
    por_periodo: dict[str, list[Registro]] = defaultdict(list)
    for registro in registros:
        por_periodo[registro.periodo].append(registro)

    bloques = []
    for periodo in sorted(por_periodo):
        registros_periodo = sorted(por_periodo[periodo], key=lambda item: item.fecha_hora)
        resumen_total = resumir_periodo(registros_periodo)
        resumenes_jornada = agrupar_resumenes_jornada(registros_periodo)

        filas_resumen = []
        for item in resumenes_jornada:
            filas_resumen.append(
                "<tr>"
                f"<td>{escape(item.jornada)}</td>"
                f"<td>{escape(item.turno)}</td>"
                f"<td>{escape(item.horario)}</td>"
                f"<td>{minutos_a_texto(item.minutos_dentro)}</td>"
                f"<td>{minutos_a_texto(item.minutos_fuera)}</td>"
                f"<td>{minutos_a_texto(item.minutos_permitidos)}</td>"
                f"<td>{minutos_a_texto(item.minutos_exceso)}</td>"
                f"<td>{escape(item.estado)}</td>"
                "</tr>"
            )

        filas_movimientos = []
        for mov in registros_periodo:
            filas_movimientos.append(
                "<tr>"
                f"<td>{escape(mov.fecha)}</td>"
                f"<td>{escape(mov.hora)}</td>"
                f"<td>{escape(mov.tipo.capitalize())}</td>"
                f"<td>{escape(mov.turno)}</td>"
                f"<td>{escape(mov.jornada)}</td>"
                f"<td>{escape(mov.detalle)}</td>"
                "</tr>"
            )

        bloques.append(
            f"""
  <section>
    <h2>Periodo {escape(periodo)}</h2>
    <p class="totales">
      Turnos: {resumen_total['turnos']} |
      Programado: {minutos_a_texto(resumen_total['programados'])} |
      Dentro: {minutos_a_texto(resumen_total['dentro'])} |
      Fuera: {minutos_a_texto(resumen_total['fuera'])} |
      Permitido: {minutos_a_texto(resumen_total['permitidos'])} |
      Exceso: {minutos_a_texto(resumen_total['exceso'])}
    </p>
    <h3>Resumen por turno</h3>
    <table>
      <thead>
        <tr>
          <th>Jornada</th>
          <th>Turno</th>
          <th>Horario</th>
          <th>Dentro</th>
          <th>Fuera</th>
          <th>Permitido</th>
          <th>Exceso</th>
          <th>Estado</th>
        </tr>
      </thead>
      <tbody>
        {''.join(filas_resumen) if filas_resumen else '<tr><td colspan="8">No hay turnos calculables en este periodo.</td></tr>'}
      </tbody>
    </table>
    <h3>Movimientos registrados</h3>
    <table>
      <thead>
        <tr>
          <th>Fecha</th>
          <th>Hora</th>
          <th>Tipo</th>
          <th>Turno</th>
          <th>Jornada</th>
          <th>Detalle</th>
        </tr>
      </thead>
      <tbody>
        {''.join(filas_movimientos) if filas_movimientos else '<tr><td colspan="6">No hay movimientos.</td></tr>'}
      </tbody>
    </table>
  </section>
"""
        )

    contenido = f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>Reporte de registros</title>
  <style>
    body {{
      font-family: Arial, sans-serif;
      margin: 32px;
      color: #1f2937;
    }}
    h1, h2, h3 {{
      margin-bottom: 8px;
    }}
    section {{
      margin-top: 28px;
    }}
    p {{
      color: #4b5563;
    }}
    .totales {{
      font-weight: bold;
      color: #111827;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin-top: 12px;
      margin-bottom: 24px;
    }}
    th, td {{
      border: 1px solid #d1d5db;
      padding: 8px;
      text-align: left;
      vertical-align: top;
      font-size: 14px;
    }}
    th {{
      background: #e5eef7;
    }}
  </style>
</head>
<body>
  <h1>Reporte de entradas y salidas</h1>
  <p>Generado en hora de Colombia: {ahora_colombia().strftime("%Y-%m-%d %H:%M:%S")}</p>
  {''.join(bloques) if bloques else '<p>No hay registros todavía.</p>'}
</body>
</html>
"""
    destino.write_text(contenido, encoding="utf-8")


class AppRegistro:
    def __init__(self, root):
        cargar_ui_escritorio()
        self.root = root
        self.root.title("Control personal de entradas y salidas")
        self.root.geometry("1180x860")
        self.root.minsize(1080, 760)

        asegurar_archivo()
        self.turno_rapido_var = tk.StringVar(value=turno_actual_por_hora())
        self.estado_var = tk.StringVar(value="Listo para registrar.")
        self.periodo_var = tk.StringVar(value=f"Periodo actual: {calcular_periodo(ahora_colombia())}")
        self.resumen_var = tk.StringVar(value="")
        self.manual_fecha_var = tk.StringVar(value=ahora_colombia().strftime("%Y-%m-%d"))
        self.manual_hora_var = tk.StringVar(value="")
        self.manual_tipo_var = tk.StringVar(value="entrada")
        self.manual_turno_var = tk.StringVar(value=turno_actual_por_hora())
        self.manual_detalle_var = tk.StringVar(value="")
        self.filtro_periodo_var = tk.StringVar(value=TODOS_LOS_PERIODOS)

        self._construir_interfaz()
        self._cargar_tablas()

    def _construir_interfaz(self):
        marco = ttk.Frame(self.root, padding=16)
        marco.pack(fill="both", expand=True)

        ttk.Label(
            marco,
            text="Registro personal del hospital",
            font=("Helvetica", 20, "bold"),
        ).pack(anchor="w")

        ttk.Label(
            marco,
            text="Guarda tus movimientos por turno y calcula tiempo dentro, fuera y exceso del periodo.",
        ).pack(anchor="w", pady=(4, 16))

        rapido_frame = ttk.LabelFrame(marco, text="Registro rapido del momento", padding=12)
        rapido_frame.pack(fill="x", pady=(0, 12))

        ttk.Label(rapido_frame, text="Turno actual").pack(side="left", padx=(0, 8))
        ttk.Combobox(
            rapido_frame,
            textvariable=self.turno_rapido_var,
            values=("12h dia", "12h noche", "5h manana", "6h manana", "6h tarde"),
            state="readonly",
            width=14,
        ).pack(side="left", padx=(0, 14))

        ttk.Button(rapido_frame, text="Registrar entrada", command=lambda: self.registrar("entrada")).pack(
            side="left", padx=(0, 8)
        )
        ttk.Button(rapido_frame, text="Registrar salida", command=lambda: self.registrar("salida")).pack(
            side="left", padx=(0, 8)
        )
        ttk.Button(rapido_frame, text="Registrar libre", command=lambda: self.registrar("libre")).pack(
            side="left", padx=(0, 8)
        )
        ttk.Button(rapido_frame, text="Exportar HTML para imprimir", command=self.exportar_como_html).pack(
            side="right"
        )
        ttk.Button(rapido_frame, text="Exportar CSV", command=self.exportar_como_csv).pack(
            side="right", padx=(0, 8)
        )

        manual_frame = ttk.LabelFrame(marco, text="Registro manual de fecha pasada", padding=12)
        manual_frame.pack(fill="x", pady=(0, 12))

        ttk.Label(manual_frame, text="Fecha").grid(row=0, column=0, sticky="w", padx=(0, 8))
        ttk.Entry(manual_frame, textvariable=self.manual_fecha_var, width=14).grid(
            row=1, column=0, sticky="w", padx=(0, 8)
        )

        ttk.Label(manual_frame, text="Hora (HH:MM)").grid(row=0, column=1, sticky="w", padx=(0, 8))
        ttk.Entry(manual_frame, textvariable=self.manual_hora_var, width=12).grid(
            row=1, column=1, sticky="w", padx=(0, 8)
        )

        ttk.Label(manual_frame, text="Tipo").grid(row=0, column=2, sticky="w", padx=(0, 8))
        ttk.Combobox(
            manual_frame,
            textvariable=self.manual_tipo_var,
            values=("entrada", "salida", "libre"),
            state="readonly",
            width=10,
        ).grid(row=1, column=2, sticky="w", padx=(0, 8))

        ttk.Label(manual_frame, text="Turno").grid(row=0, column=3, sticky="w", padx=(0, 8))
        ttk.Combobox(
            manual_frame,
            textvariable=self.manual_turno_var,
            values=("12h dia", "12h noche", "5h manana", "6h manana", "6h tarde", "libre"),
            state="readonly",
            width=12,
        ).grid(row=1, column=3, sticky="w", padx=(0, 8))

        ttk.Label(manual_frame, text="Detalle").grid(row=0, column=4, sticky="w", padx=(0, 8))
        ttk.Entry(manual_frame, textvariable=self.manual_detalle_var, width=28).grid(
            row=1, column=4, sticky="ew", padx=(0, 8)
        )

        ttk.Button(manual_frame, text="Guardar manual", command=self.registrar_manual).grid(
            row=1, column=5, sticky="e"
        )
        manual_frame.columnconfigure(4, weight=1)

        ttk.Label(marco, textvariable=self.estado_var, foreground="#1d4ed8").pack(anchor="w", pady=(0, 8))
        ttk.Label(marco, textvariable=self.periodo_var, foreground="#374151").pack(anchor="w", pady=(0, 8))

        filtro_frame = ttk.Frame(marco)
        filtro_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(filtro_frame, text="Ver periodo").pack(side="left", padx=(0, 8))
        self.filtro_combo = ttk.Combobox(
            filtro_frame,
            textvariable=self.filtro_periodo_var,
            state="readonly",
            width=32,
        )
        self.filtro_combo.pack(side="left")
        self.filtro_combo.bind("<<ComboboxSelected>>", lambda event: self._cargar_tablas())

        resumen_frame = ttk.LabelFrame(marco, text="Resumen del periodo filtrado", padding=12)
        resumen_frame.pack(fill="x", pady=(0, 12))
        ttk.Label(
            resumen_frame,
            textvariable=self.resumen_var,
            foreground="#111827",
            justify="left",
        ).pack(anchor="w")

        ttk.Label(marco, text="Resumen por jornada").pack(anchor="w")
        columnas_resumen = (
            "jornada",
            "turno",
            "horario",
            "dentro",
            "fuera",
            "permitido",
            "exceso",
            "estado",
        )
        self.tabla_resumen = ttk.Treeview(marco, columns=columnas_resumen, show="headings", height=8)
        for columna, titulo, ancho in (
            ("jornada", "Jornada", 100),
            ("turno", "Turno", 100),
            ("horario", "Horario", 260),
            ("dentro", "Dentro", 80),
            ("fuera", "Fuera", 80),
            ("permitido", "Permitido", 80),
            ("exceso", "Exceso", 80),
            ("estado", "Estado", 180),
        ):
            self.tabla_resumen.heading(columna, text=titulo)
            self.tabla_resumen.column(columna, width=ancho, anchor="center")
        self.tabla_resumen.pack(fill="x", pady=(4, 14))

        ttk.Label(marco, text="Movimientos registrados").pack(anchor="w")
        columnas = ("fecha", "hora", "tipo", "turno", "jornada", "detalle", "periodo")
        self.tabla = ttk.Treeview(marco, columns=columnas, show="headings", height=14)
        for columna, titulo, ancho, anchor in (
            ("fecha", "Fecha", 95, "center"),
            ("hora", "Hora", 80, "center"),
            ("tipo", "Tipo", 80, "center"),
            ("turno", "Turno", 100, "center"),
            ("jornada", "Jornada", 95, "center"),
            ("detalle", "Detalle", 280, "w"),
            ("periodo", "Periodo", 190, "center"),
        ):
            self.tabla.heading(columna, text=titulo)
            self.tabla.column(columna, width=ancho, anchor=anchor)

        tabla_frame = ttk.Frame(marco)
        tabla_frame.pack(fill="both", expand=True)
        self.tabla.pack(in_=tabla_frame, side="left", fill="both", expand=True)
        scroll = ttk.Scrollbar(tabla_frame, orient="vertical", command=self.tabla.yview)
        self.tabla.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")

    def registrar(self, tipo: str):
        turno = "libre" if tipo == "libre" else self.turno_rapido_var.get()
        detalle = "Dia libre" if tipo == "libre" else ""
        registro = guardar_registro(tipo=tipo, turno=turno, detalle=detalle)
        self.periodo_var.set(f"Periodo actual: {calcular_periodo(ahora_colombia())}")
        self.estado_var.set(
            f"Registrado: {registro.tipo.capitalize()} el {registro.fecha} a las {formatear_hora_visible(registro.hora)} en turno {registro.turno}."
        )
        self._cargar_tablas()

    def _registros_filtrados(self, registros: list[Registro]) -> list[Registro]:
        periodo_seleccionado = self.filtro_periodo_var.get()
        if periodo_seleccionado == TODOS_LOS_PERIODOS:
            return registros
        return [registro for registro in registros if registro.periodo == periodo_seleccionado]

    def _cargar_tablas(self):
        for item in self.tabla.get_children():
            self.tabla.delete(item)
        for item in self.tabla_resumen.get_children():
            self.tabla_resumen.delete(item)

        registros = leer_registros()
        self.filtro_combo["values"] = periodos_disponibles(registros)
        if self.filtro_periodo_var.get() not in self.filtro_combo["values"]:
            self.filtro_periodo_var.set(TODOS_LOS_PERIODOS)

        registros_filtrados = self._registros_filtrados(registros)
        resumen_total = resumir_periodo(registros_filtrados)
        self.resumen_var.set(
            " | ".join(
                [
                    f"Turnos calculados: {resumen_total['turnos']}",
                    f"Tiempo programado: {minutos_a_texto(resumen_total['programados'])}",
                    f"Tiempo dentro: {minutos_a_texto(resumen_total['dentro'])}",
                    f"Tiempo fuera: {minutos_a_texto(resumen_total['fuera'])}",
                    f"Salida permitida: {minutos_a_texto(resumen_total['permitidos'])}",
                    f"Exceso a revisar: {minutos_a_texto(resumen_total['exceso'])}",
                    f"Registros sin turno: {resumen_total['sin_turno']}",
                ]
            )
        )

        for item in agrupar_resumenes_jornada(registros_filtrados):
            self.tabla_resumen.insert(
                "",
                "end",
                values=(
                    item.jornada,
                    item.turno,
                    item.horario,
                    minutos_a_texto(item.minutos_dentro),
                    minutos_a_texto(item.minutos_fuera),
                    minutos_a_texto(item.minutos_permitidos),
                    minutos_a_texto(item.minutos_exceso),
                    item.estado,
                ),
            )

        for registro in registros_filtrados:
            self.tabla.insert(
                "",
                "end",
                values=(
                    registro.fecha,
                    formatear_hora_visible(registro.hora),
                    registro.tipo.capitalize(),
                    registro.turno,
                    registro.jornada,
                    registro.detalle,
                    registro.periodo,
                ),
            )

    def registrar_manual(self):
        fecha_texto = self.manual_fecha_var.get().strip()
        hora_texto = self.manual_hora_var.get().strip()
        tipo = self.manual_tipo_var.get().strip().lower()
        turno = self.manual_turno_var.get().strip().lower()
        detalle = self.manual_detalle_var.get().strip()

        if tipo not in {"entrada", "salida", "libre"}:
            messagebox.showerror("Tipo invalido", "El tipo debe ser entrada, salida o libre.")
            return
        if turno not in TURNOS:
            messagebox.showerror("Turno invalido", "Escoge un turno valido.")
            return
        if tipo == "libre":
            turno = "libre"
            if not detalle:
                detalle = "Dia libre"
        elif turno == "libre":
            messagebox.showerror("Turno invalido", "Entrada o salida no pueden quedar con turno libre.")
            return

        try:
            fecha_hora = construir_fecha_hora_manual(fecha_texto, hora_texto)
        except ValueError:
            messagebox.showerror(
                "Fecha u hora invalida",
                "Usa el formato fecha AAAA-MM-DD y hora HH:MM.",
            )
            return

        registro = guardar_registro_en_fecha(tipo=tipo, fecha_hora=fecha_hora, turno=turno, detalle=detalle)
        self.periodo_var.set(f"Periodo actual: {calcular_periodo(ahora_colombia())}")
        self.estado_var.set(
            f"Registrado manual: {registro.tipo.capitalize()} del {registro.fecha} a las {formatear_hora_visible(registro.hora)} en turno {registro.turno}."
        )
        self.manual_detalle_var.set("")
        self._cargar_tablas()

    def exportar_como_csv(self):
        registros = self._registros_filtrados(leer_registros())
        destino = filedialog.asksaveasfilename(
            title="Guardar historial como CSV",
            defaultextension=".csv",
            filetypes=[("Archivo CSV", "*.csv")],
            initialfile="historial_organizado.csv",
        )
        if not destino:
            return

        exportar_csv(Path(destino), registros)
        messagebox.showinfo("Exportacion lista", "Se guardo el historial filtrado en formato CSV.")

    def exportar_como_html(self):
        registros = self._registros_filtrados(leer_registros())
        destino = filedialog.asksaveasfilename(
            title="Guardar reporte para imprimir",
            defaultextension=".html",
            filetypes=[("Archivo HTML", "*.html")],
            initialfile="reporte_para_imprimir.html",
        )
        if not destino:
            return

        exportar_html(Path(destino), registros)
        messagebox.showinfo(
            "Reporte listo",
            "Se guardo el reporte HTML con resumen de turnos y movimientos.",
        )


def main():
    cargar_ui_escritorio()
    root = tk.Tk()
    style = ttk.Style()
    if "clam" in style.theme_names():
        style.theme_use("clam")
    AppRegistro(root)
    root.mainloop()


if __name__ == "__main__":
    main()
